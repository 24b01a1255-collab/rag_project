import docx

from langchain_core.documents import Document


def load_docx(uploaded_file):

    doc = docx.Document(uploaded_file)

    full_text = []

    for para in doc.paragraphs:

        full_text.append(para.text)

    text = "\n".join(full_text)

    documents = [

        Document(

            page_content=text,

            metadata={

                "source": uploaded_file.name,

                "page": 1
            }
        )
    ]

    return documents